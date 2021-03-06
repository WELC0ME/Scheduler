import sys
import sqlite3
from PyQt5.QtWidgets import QApplication, QMainWindow, QHeaderView, \
    QTableWidgetItem, QCheckBox, QSpinBox
from PyQt5.QtGui import QIcon
from PyQt5 import uic
from PyQt5.QtCore import Qt
from config import *
import config
import datetime
import os
import shutil
from docx import Document


class Dialog(QMainWindow):

    def __init__(self, name, db_table, deps):
        super().__init__()

        uic.loadUi('interfaces/' + name + '.ui', self)
        self.setWindowTitle(name.capitalize())
        self.setStyleSheet(STYLESHEET)
        self.btn_save.clicked.connect(self.save)
        self.btn_close.clicked.connect(self.close)
        self.db_table = db_table
        self.deps = deps
        self.copy = None

        con = sqlite3.connect(DATABASE_NAME)
        cur = con.cursor()
        config.OBJECTS[name] = [self.new_object(i) for i in
                                cur.execute('SELECT * FROM ' + db_table)]
        con.close()

    def new_object(self, params):
        obj = Obj(params[0])
        for i, attr in enumerate(self.deps.keys()):
            try:
                setattr(obj, attr, eval(str(params[1 + i])))
            except Exception:
                setattr(obj, attr, eval('"' + str(params[1 + i]) + '"'))
        return obj

    def show(self):
        self.statusBar().showMessage("")
        self.manage('set')
        super().show()

    def save(self):
        if self.manage('get'):
            if config.CHOSEN_KEY[1] == -1:
                config.OBJECTS[config.CHOSEN_KEY[0]].append(self.copy.copy())
            else:
                config.OBJECTS[config.CHOSEN_KEY[0]][
                    config.CHOSEN_KEY[1]] = self.copy.copy()
            self.copy = None
            self.hide()
            windows['menu'].show()

    def close(self):
        self.copy = None
        self.hide()
        windows['menu'].show()

    def manage(self, state):
        try:
            self.copy = config.CHOSEN_KEY[
                2].copy() if not self.copy else self.copy
            for attr in self.deps.keys():
                getattr(self, self.deps[attr][:2] + '_' + state)(
                    getattr(self, self.deps[attr]), self.copy, attr)
            return True
        except AssertionError as e:
            self.statusBar().showMessage("Error: %s" % e)

    @staticmethod
    def cb_get(cb, obj, attr):
        setattr(obj, attr, cb.currentText())

    @staticmethod
    def cb_set(cb, obj, attr):
        if not hasattr(obj, attr):
            setattr(obj, attr, 1)
        cb.setCurrentIndex(int(getattr(obj, attr)) - 1)

    @staticmethod
    def tw_get(tw, obj, attr):
        keys = {
            'QCheckBox': 'isChecked',
            'QSpinBox': 'value'
        }
        out = getattr(obj, attr)
        values = list(getattr(obj, attr)['values'].keys())
        columns = [i for i in range(tw.columnCount())] if attr == 'days' else [1]
        for i in range(tw.rowCount()):
            for k in columns:
                res = getattr(tw.cellWidget(i, k), keys[
                    str(type(tw.cellWidget(i, k))).split('.')[-1][:-2]])()
                if isinstance(res, bool):
                    res = 1 if res else 0
                shift = k if attr == 'days' else 0
                coeff = tw.columnCount() if attr == 'days' else 1
                out['values'][values[i * coeff + shift]] = res
        setattr(obj, attr, out)

    @staticmethod
    def tw_set(tw, obj, attr):
        set_stretch(tw)

        gen = lambda name: {i.id_: 0 for i in config.OBJECTS[name]}
        keys = {
            'subjects': {'values': gen('subject'), 'type': 'number', 'table': 'subject'},
            'asubjects': {'values': gen('subject'), 'type': 'bool', 'table': 'subject'},
            'classes': {'values': gen('class'), 'type': 'bool', 'table': 'class'},
            'rooms': {'values': gen('room'), 'type': 'bool', 'table': 'room'},
            'days': {'values': {i: 0 for i in range(48)}, 'type': 'bool', 'table': '.'},
        }
        res = keys[attr]
        for i in res['values'].keys():
            try:
                res['values'][i] = getattr(obj, attr)['values'][i]
            except Exception:
                continue
        setattr(obj, attr, res)

        keys = {
            'bool': [QCheckBox, 'setChecked'],
            'number': [QSpinBox, 'setValue']
        }

        tw.setRowCount(len(res['values'].keys()) if attr != 'days' else 6)
        values = list(res['values'].keys())
        if attr != 'days':
            for i in range(tw.rowCount()):
                tw.setItem(i, 0, QTableWidgetItem(str(
                    [j for j in config.OBJECTS[res['table']] if
                     j.id_ == values[i]][0].name)))
                tw.setCellWidget(i, 1, keys[res['type']][0]())
                getattr(tw.cellWidget(i, 1), keys[res['type']][1])(
                    res['values'][values[i]])
        else:
            for i in range(tw.rowCount()):
                for k in range(tw.columnCount()):
                    tw.setCellWidget(i, k, keys[res['type']][0]())
                    getattr(tw.cellWidget(i, k), keys[res['type']][1])(
                        res['values'][values[i * tw.columnCount() + k]])

    @staticmethod
    def le_get(le, obj, attr):
        assert le.text(), 'line edit unfilled'
        setattr(obj, attr, le.text())

    @staticmethod
    def le_set(le, obj, attr):
        if not hasattr(obj, attr):
            setattr(obj, attr, '')
        le.setText(str(getattr(obj, attr)))


class Menu(QMainWindow):

    def __init__(self):
        super().__init__()
        uic.loadUi('interfaces/menu.ui', self)
        self.setStyleSheet(STYLESHEET)
        self.setWindowTitle('Scheduler')

        self.tabWidget.currentChanged.connect(self.exit_)
        self.btn_create.clicked.connect(self.create_schedule)

        for name in config.OBJECTS.keys():
            getattr(self, 'btn_add_' + name).clicked.connect(self.add)
            getattr(self, 'tw_' + name).cellDoubleClicked.connect(self.add)
            set_stretch(getattr(self, 'tw_' + name))
            getattr(self, 'btn_remove_' + name).clicked.connect(self.remove)
        self.deps = {
            0: None,
            1: 'class',
            2: 'teacher',
            3: 'subject',
            4: 'room'
        }

    def show(self):
        for name in config.OBJECTS.keys():
            getattr(self, 'tw_' + name).setRowCount(len(config.OBJECTS[name]))
            column = 0
            for attr in windows[name].deps.keys():
                if windows[name].deps[attr][:2] in ['le', 'cb']:
                    for i, obj in enumerate(config.OBJECTS[name]):
                        getattr(self, 'tw_' + name).setItem(
                            i, column, QTableWidgetItem(str(getattr(obj, attr))))
                    column += 1
        super().show()

    def add(self, row):
        if isinstance(row, bool):
            ids = [i.id_ for i in config.OBJECTS[self.deps[self.tabWidget.currentIndex()]]]
            row = -1
            obj = Obj(max(ids) + 1 if ids else 0)
        else:
            obj = config.OBJECTS[self.deps[self.tabWidget.currentIndex()]][row]

        config.CHOSEN_KEY = (
            self.deps[self.tabWidget.currentIndex()], row, obj)
        self.hide()
        windows[self.deps[self.tabWidget.currentIndex()]].show()

    def remove(self):
        objects = config.OBJECTS[self.deps[self.tabWidget.currentIndex()]]
        for k in sorted(list(set([idx.row() for idx in getattr(
                self, 'tw_' + self.deps[self.tabWidget.currentIndex()]
        ).selectedIndexes()])))[::-1]:
            objects.remove(objects[k])
        self.show()

    def exit_(self):
        if self.tabWidget.currentIndex() == 5:
            con = sqlite3.connect(DATABASE_NAME)
            cur = con.cursor()
            for window in windows.keys():
                if not isinstance(windows[window], Dialog):
                    continue
                cur.execute("DELETE FROM " + windows[window].db_table)
                for obj in config.OBJECTS[window]:
                    cur.execute('INSERT INTO ' + windows[window].db_table + '(' + ', '.join(
                        ['id', *[str(i) for i in windows[window].deps.keys()]]) + ') VALUES (' + ', '.join([
                        '"' + str(getattr(obj, i)) + '"' for i in [
                            str(k) for k in ['id_', *windows[window].deps.keys()]]]) + ')')
            con.commit()
            con.close()
            sys.exit(1)

    def create_schedule(self):
        table = []
        for day in range(6 * 8):
            table.append([])
            for class_ in config.OBJECTS['class']:
                table[day].append([])
                for teacher_ in config.OBJECTS['teacher']:
                    if class_.id_ in teacher_.classes['values'].keys():
                        if teacher_.days['values'][day]:
                            for subject_ in teacher_.asubjects['values'].keys():
                                for room_ in teacher_.rooms['values'].keys():
                                    table[day][-1].append({
                                        'subject': subject_,
                                        'room': room_,
                                        'teacher': teacher_.id_,
                                    })
        table.append([])
        for class_ in config.OBJECTS['class']:
            table[-1].append({})
            for subject in list(class_.subjects['values'].keys()):
                table[-1][-1][subject] = class_.subjects['values'][subject]
        config.SOLVED = False
        config.RESULT = []
        try:
            self.solve(table, 0, 0)
        except Exception:
            config.SOLVED = False
            config.RESULT = []
        if config.SOLVED:
            try:
                shutil.rmtree('Schedule', ignore_errors=True)
                os.mkdir('Schedule')
                os.mkdir('Schedule/classes')
                os.mkdir('Schedule/teachers')
                _classes = {i.id_: i for i in config.OBJECTS['class']}
                _teachers = {i.id_: i for i in config.OBJECTS['teacher']}
                _rooms = {i.id_: i for i in config.OBJECTS['room']}
                _subjects = {i.id_: i for i in config.OBJECTS['subject']}
                for class_ in range(len(config.RESULT[0])):
                    document = Document()
                    document.add_heading(str(_classes[class_].name), 0)
                    table = document.add_table(rows=7, cols=9)
                    week_days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday',
                                 'Friday', 'Saturday']
                    for i in range(1, 9):
                        table.rows[0].cells[i].text = str(i - 1)
                    for i in range(6):
                        table.rows[1 + i].cells[0].text = week_days[i]
                    for day in range(len(config.RESULT) - 1):
                        if len(config.RESULT[day][class_]) == 1:
                            teacher_ = _teachers[config.RESULT[day][class_][0]['teacher']]
                            info = ' '.join([
                                str(teacher_.surname),
                                str(teacher_.name),
                                str(teacher_.patronymic) + '\n',
                                str(_rooms[config.RESULT[day][class_][0]['room']].name),
                                str(_subjects[config.RESULT[day][class_][0]['subject']].name),
                            ])
                            table.rows[day // 8 + 1].cells[day % 8 + 1].text = info
                    document.save('Schedule/classes/' + str(_classes[class_].name) + '.docx')

                for teacher_ in [i.id_ for i in config.OBJECTS['teacher']]:
                    document = Document()
                    initials = ' '.join([
                        str(_teachers[teacher_].surname),
                        str(_teachers[teacher_].name),
                        str(_teachers[teacher_].patronymic),
                    ])
                    document.add_heading(initials, 0)
                    table = document.add_table(rows=7, cols=9)
                    week_days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday',
                                 'Friday', 'Saturday']
                    for i in range(1, 9):
                        table.rows[0].cells[i].text = str(i - 1)
                    for i in range(6):
                        table.rows[1 + i].cells[0].text = week_days[i]
                    for day in range(len(config.RESULT) - 1):
                        for class_ in range(len(config.RESULT[0])):
                            if len(config.RESULT[day][class_]) == 1:
                                if config.RESULT[day][class_][0]['teacher'] == teacher_:
                                    info = ' '.join([
                                        str(_classes[class_].name),
                                        str(_rooms[config.RESULT[day][class_][0]['room']].name),
                                        str(_subjects[config.RESULT[day][class_][0]['subject']].name),
                                    ])
                                    table.rows[day // 8 + 1].cells[day % 8 + 1].text = info
                        document.save('Schedule/teachers/' + initials + '.docx')
                message = 'Success'
            except:
                message = 'Fail'
        else:
            message = 'Fail'
        time = '[' + str(datetime.datetime.now()).split('.')[0].split()[1] + '] '
        self.log.setPlainText(self.log.toPlainText() + time + message + '\n')

    def solve(self, table, lesson, class_):
        if config.SOLVED:
            return
        if class_ == len(table[lesson]):
            class_ = 0
            lesson += 1
        if lesson + 1 > 6 * 8:
            if all([all([j == 0 for j in k.values()]) for k in table[-1]]):
                config.RESULT = self.copy_table(table)
                config.SOLVED = True
            return

        for i in range(len(table[-1])):
            targets = table[-1][i]
            for j in targets.keys():
                lesson_counter = 0
                for k in range(len(table) - 1):
                    lesson_counter += 1 if len(
                        [h for h in table[k][i] if
                         h['subject'] == j]) >= 1 else 0
                if lesson_counter < targets[j]:
                    return
        continuing = False
        for i in table[lesson][class_]:

            if table[-1][class_][i['subject']] > 0:
                continuing = True
                out = self.copy_table(table)
                out[lesson][class_] = [i]
                out[-1][class_][i['subject']] -= 1
                for k in range(class_, len(table[lesson])):
                    table[lesson][k] = [j for j in table[lesson][k]
                                        if (j['teacher'] != i['teacher'] and
                                            j['room'] != i['room'])]
                for k in range(lesson, len(table) - 1):
                    table[k][class_] = [j for j in table[k][class_]
                                        if (j['subject'] != i['subject'] or
                                            j['teacher'] == i['teacher'])]
                self.solve(out, lesson, class_ + 1)

        if not continuing:
            out = self.copy_table(table)
            out[lesson][class_] = []
            self.solve(out, lesson, class_ + 1)

    def copy_table(self, table):
        return [[table[i][k].copy() for k in range(len(table[i]))] for i in
                range(len(table))]


def set_stretch(table):
    for i in range(table.columnCount()):
        table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)
    for i in range(table.rowCount()):
        table.verticalHeader().setSectionResizeMode(i, QHeaderView.Stretch)


def except_hook(cls, exception, traceback):
    sys.__excepthook__(cls, exception, traceback)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon('data/icon.png'))

    windows = {
        'class': Dialog('class', 'Classes', {
            'name': 'le_name',
            'number': 'cb_number',
            'subjects': 'tw_subjects'
        }),
        'teacher': Dialog('teacher', 'Teachers', {
            'surname': 'le_surname',
            'name': 'le_name',
            'patronymic': 'le_patronymic',
            'classes': 'tw_classes',
            'asubjects': 'tw_asubjects',
            'rooms': 'tw_rooms',
            'days': 'tw_days'
        }),
        'subject': Dialog('subject', 'Subjects', {
            'name': 'le_name',
        }),
        'room': Dialog('room', 'Rooms', {
            'name': 'le_name'
        }),
        'menu': Menu(),
    }
    windows['menu'].show()
    sys.excepthook = except_hook
    sys.exit(app.exec())
